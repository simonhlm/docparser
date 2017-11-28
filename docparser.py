import docx, json, logging, re
from pprint import pprint

logging.basicConfig(level='DEBUG')

delimter = ["：", "。","；",":","#"]  

para_list = {
    "入院情况":"entry_info",
    "入院诊断":"entry_diagnose",
    "诊疗经过":"diagnose_process",
    "出院诊断":"leave_diagnose",
    "出院情况":"leave_info",
    "出院医嘱":"leave_advise"
}

def read_docx(docName):
    doc = docx.Document(docName)
    return doc

def parse_docx(doc, doc_dict):
    logging.info('entry_info')

    for p in doc.paragraphs:
        p_dict = {}
        sentenses = []
        single_sentence = []
        p_list = []
        text = re.sub('\d+\.(?!\d)', "#", p.text)
        logging.info(text)

        for chars in text:
            if chars not in delimter:
                single_sentence.append(chars)
                #logging.debug(single_sentence)
            else:
                sentenses.append(''.join(single_sentence))
                #logging.debug(sentenses[-1])
                if chars == "：":
                    p_list = []
                    p_dict[sentenses[-1]] = p_list         
                elif len(sentenses[-1]) > 0:
                    p_list.append(sentenses[-1])
                single_sentence = []

        for key in para_list:
            if p.text.find(key) == 0:
                doc_dict[key] = p_dict

    return doc_dict

def parse_table(doc, doc_dict):
    tables = doc.tables
    basic_info_dict = {}
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                cell_c = re.split("[:：]",cell.text)
                basic_info_dict[cell_c[0]] = cell_c[1]
    doc_dict['基本信息'] = basic_info_dict
    return doc_dict


if __name__ == '__main__':
    doc_dict = {}
    doc = read_docx('马清喜出院记录.docx')

    doc_dict = parse_table(doc, doc_dict)
    doc_dict = parse_docx(doc, doc_dict)

    #print(json.dumps(output, ensure_ascii=False, indent =4))

    with open('test.json','w', encoding = 'utf-8') as f:
        json.dump(doc_dict, f, ensure_ascii=False, indent =4 )